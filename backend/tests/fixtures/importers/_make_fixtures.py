"""產生 parser 測試用的 fixture 檔。

執行方式：`uv run python tests/fixtures/importers/_make_fixtures.py`
純文字 fixture（csv / md / html / ics）直接寫在這裡，二進位 fixture
（xlsx / pdf / docx）由對應套件產生，以免把二進位檔塞進 git diff 難以審閱。
"""

from __future__ import annotations

from pathlib import Path

HERE = Path(__file__).parent

CSV = """title,start,end,note
晨跑,2026-09-08T07:00:00Z,2026-09-08T07:40:00Z,輕鬆配速
重訓,2026-09-10T19:00:00Z,2026-09-10T20:00:00Z,上肢
"""

CSV_NO_DATE = """name,note
run,easy pace
lift,upper body
"""

MD = """# 訓練筆記

目前 5K 大約 38 分，想在年底前跑進 30 分。

## 現況

每週能跑三次，平日晚上與週六早上。

## 限制

左膝舊傷，不能連續兩天高強度。
"""

HTML = """<html><head><title>訓練筆記</title></head><body>
<h1>訓練筆記</h1>
<p>目前 5K 大約 38 分。</p>
<h2>現況</h2>
<p>每週能跑三次。</p>
</body></html>
"""

ICS = """BEGIN:VCALENDAR
VERSION:2.0
PRODID:-//guru-core//test//EN
BEGIN:VEVENT
UID:evt-1@example.com
DTSTART:20260908T090000Z
DTEND:20260908T100000Z
SUMMARY:週會
LOCATION:會議室 A
END:VEVENT
BEGIN:VEVENT
UID:evt-2@example.com
DTSTART;VALUE=DATE:20260912
DTEND;VALUE=DATE:20260913
SUMMARY:出差
END:VEVENT
END:VCALENDAR
"""


def _write_text() -> None:
    (HERE / "sample.csv").write_text(CSV, encoding="utf-8")
    (HERE / "sample_no_date.csv").write_text(CSV_NO_DATE, encoding="utf-8")
    (HERE / "sample.md").write_text(MD, encoding="utf-8")
    (HERE / "sample.html").write_text(HTML, encoding="utf-8")
    (HERE / "sample.ics").write_text(ICS, encoding="utf-8")
    (HERE / "empty.csv").write_text("", encoding="utf-8")
    (HERE / "empty.md").write_text("", encoding="utf-8")


def _write_xlsx() -> None:
    from openpyxl import Workbook

    wb = Workbook()
    ws = wb.active
    assert ws is not None
    ws.title = "訓練"
    ws.append(["title", "start", "end", "note"])
    ws.append(["晨跑", "2026-09-08T07:00:00Z", "2026-09-08T07:40:00Z", "輕鬆配速"])
    ws.append(["重訓", "2026-09-10T19:00:00Z", "2026-09-10T20:00:00Z", "上肢"])
    wb.save(HERE / "sample.xlsx")


def _write_pdf() -> None:
    from pypdf import PdfWriter

    writer = PdfWriter()
    writer.add_blank_page(width=595, height=842)
    writer.add_blank_page(width=595, height=842)
    with (HERE / "sample.pdf").open("wb") as fh:
        writer.write(fh)


def _write_docx() -> None:
    from docx import Document as DocxDocument

    doc = DocxDocument()
    doc.add_heading("訓練筆記", level=1)
    doc.add_paragraph("目前 5K 大約 38 分。")
    doc.add_heading("現況", level=2)
    doc.add_paragraph("每週能跑三次。")
    doc.save(str(HERE / "sample.docx"))


if __name__ == "__main__":
    _write_text()
    _write_xlsx()
    _write_pdf()
    _write_docx()
    print("fixtures written to", HERE)
