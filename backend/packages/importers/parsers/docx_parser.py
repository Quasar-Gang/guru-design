"""DOCX parser: collect non-empty paragraphs, starting a new chunk at every heading."""

from __future__ import annotations

import io

from docx import Document as load_docx

from packages.importers.document import Document, TextChunk
from packages.importers.ports import RawBlob


def _is_heading(style_name: str) -> bool:
    lowered = style_name.strip().lower()
    return lowered.startswith("heading") or lowered == "title"


class DocxParser:
    """Parse .docx; the text of a heading paragraph becomes the chunk's section."""

    def supports(self, fmt: str) -> bool:
        return fmt == "docx"

    def parse(self, blob: RawBlob) -> Document:
        if not blob.data:
            return Document()

        chunks: list[TextChunk] = []
        section: str | None = None
        body: list[str] = []

        def flush() -> None:
            joined = "\n\n".join(part for part in body if part)
            if section is not None or joined:
                chunks.append(TextChunk(text=joined, section=section, order=len(chunks)))
            body.clear()

        for paragraph in load_docx(io.BytesIO(blob.data)).paragraphs:
            content = paragraph.text.strip()
            if not content:
                continue
            style = paragraph.style
            style_name = "" if style is None or style.name is None else style.name
            if _is_heading(style_name):
                flush()
                section = content
                continue
            body.append(content)
        flush()
        return Document(text_chunks=chunks)
